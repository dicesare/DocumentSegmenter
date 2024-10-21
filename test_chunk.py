import os
import csv
import json
from docx import Document
from datetime import datetime
import pymupdf4llm as pymu  # Library to handle PDF extraction


class Segment:
    """
    Represents a segment of extracted text (title, subtitle, paragraph).

    :param text: The segment's text content
    :param segment_type: Type of the segment (title, subtitle, paragraph)
    :param importance: Relative importance of the segment (default is 1.0)
    :return: A Segment object
    """

    def __init__(self, text, segment_type, importance=1.0):
        self.text = text
        self.segment_type = segment_type
        self.importance = importance

    def __repr__(self):
        return f"Segment(type={self.segment_type}, importance={self.importance}, text={self.text[:30]}...)"


class DocumentSegmenter:
    """
    A class that segments and extracts titles, subtitles, and paragraphs from DOCX, PDF, and TXT documents.

    :param file_path: The path to the file to be processed
    :param output_dir: Directory where the CSV files will be saved
    :return: A DocumentSegmenter object

    EXAMPLE USAGE:
    file_path = r'data/original/your_document_to_chunk.pdf' or document.docx  # Replace with your file path
    output_dir = 'data/chunk'  # Folder to save the CSV file; the folder will be created if it doesn't exist
    """
    def __init__(self, file_path, output_dir, save_format='csv'):

        # Validate the save format (only 'csv' or 'json' are allowed)
        if save_format.lower() not in ['csv', 'json']:
            raise ValueError(f"Unsupported save format: {save_format}")
        self.save_format = save_format.lower()

        # Ensure the directory of the file to chunk exists
        directory_path = os.path.dirname(file_path)
        self._ensure_directory_exists(directory_path)

        # Ensure the output directory exists
        self._ensure_directory_exists(output_dir)
        self.output_dir = output_dir
        
        # Check if the input file exists after creating the directory (if necessary)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file {file_path} does not exist.")
        self.file_path = file_path
        
        # Determine the file extension and initialize segments list
        self.file_extension = os.path.splitext(file_path)[-1].lower()
        self.segments = []

    def _ensure_directory_exists(self, directory_path):
        """
        Ensures that a directory exists, creating it if necessary.
        :param directory_path: The path to the directory that needs to be checked/created
        """
        if not os.path.exists(directory_path):
            os.makedirs(directory_path)
            print(f"Directory created: {directory_path}")
        else:
            print(f"Directory already exists: {directory_path}")

    def _load_document(self):
        """
        Loads the document based on its file type (DOCX, PDF, TXT) and returns its content as plain text or Markdown.

        :return: Plain text or Markdown representation of the document
        """
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"The file {self.file_path} does not exist.")

        if self.file_extension == '.docx':
            return self._load_docx()
        elif self.file_extension == '.pdf':
            return self._load_pdf()
        elif self.file_extension == '.txt':
            return self._load_txt()
        else:
            raise ValueError(f"Unsupported file format: {self.file_extension}")

    def _load_docx(self):
        """
        Loads a DOCX file and returns its content as plain text.

        :return: The plain text content of the document
        """
        doc = Document(self.file_path)
        # Convert all paragraphs into a single string joined by newlines
        return "\n".join([para.text for para in doc.paragraphs])

    def _load_pdf(self):
        """
        Uses pymupdf4llm to extract text from a PDF file and convert it to Markdown.

        :return: The converted Markdown text from the PDF
        """
        try:
            return pymu.to_markdown(self.file_path)  # Converts the PDF content to Markdown
        except Exception as e:
            print(f"Error reading the PDF file: {e}")
            return None

    def _load_txt(self):
        """
        Reads a TXT file and returns its content as plain text.

        :return: The plain text content of the TXT file
        """
        with open(self.file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _segment_document(self):
        """
        Segments the loaded document based on text styles and content type (title, subtitle, paragraph).
        The segmentation is determined differently for DOCX, PDF, and TXT.
        """
        text = self._load_document()
        if text is None:
            raise ValueError("Unable to load the document content.")

        # Call appropriate segmentation method based on file extension
        if self.file_extension == '.docx':
            self._segment_docx(text)
        elif self.file_extension == '.pdf':
            self._segment_pdf(text)
        else:
            self._segment_plain_text(text)  # Default segmentation for TXT

    def _segment_pdf(self, markdown_text):
        """
        Segments a PDF document that has been converted to Markdown by detecting titles, subtitles, and paragraphs.

        This method uses Markdown symbols (#, ##, etc.) to identify different levels of headings and segments the content accordingly.

        :param markdown_text: The Markdown text extracted from the PDF
        """
        current_segment_type = "paragraph"
        current_importance = 1.0
        current_text = ""

        # Iterate through each line of the markdown text
        for line in markdown_text.splitlines():
            line = line.strip()  # Remove leading/trailing whitespace

            if not line:  # Skip empty lines
                continue

            # Detect titles based on the Markdown heading symbols (e.g., # for title, ## for subtitle)
            if line.startswith('#'):
                # If a paragraph was being built, append it before switching segment type
                if current_text and current_segment_type == "paragraph":
                    self.segments.append(Segment(current_text.strip(), current_segment_type, current_importance))
                    current_text = ""

                # Count the number of '#' to determine heading level
                level = len(line.split(' ')[0])

                # Set segment type and importance based on heading level
                if level == 1:
                    current_segment_type = 'title'
                    current_importance = 2.0
                elif level == 2:
                    current_segment_type = 'subtitle'
                    current_importance = 1.8
                else:
                    current_segment_type = 'subtitle'  # Treat any heading level 3+ as a subtitle
                    current_importance = 1.6

                # Append the detected title/subtitle to the segments
                self.segments.append(Segment(line.strip('#').strip(), current_segment_type, current_importance))
                current_segment_type = "paragraph"  # Reset for the next paragraph
                current_importance = 1.0

            else:
                # Collect text for paragraphs
                current_text += line + " "

        # Add any remaining paragraph text to the segments
        if current_text.strip():
            self.segments.append(Segment(current_text.strip(), "paragraph", 1.0))

    def _segment_docx(self, text):
        """
        Segments a DOCX document by detecting titles, subtitles, and paragraphs based on text styles (e.g., Heading 1, Heading 2).

        This method ensures that titles and subtitles are properly identified based on their style level in DOCX.

        :return: A list of detected segments (title, subtitle, paragraph)
        """
        doc = Document(self.file_path)
        current_text = ""
        current_segment_type = "paragraph"
        current_importance = 1.0

        # Iterate through each paragraph in the DOCX file
        for para in doc.paragraphs:
            # Replace non-breaking spaces and strip excess whitespace
            line = para.text.replace('\u00A0', ' ').strip()

            if not line:
                continue  # Skip empty paragraphs

            # Detect titles and subtitles based on heading levels
            if para.style.name.startswith('Heading'):  # Title or subtitle
                if current_text and current_segment_type == 'paragraph':
                    # Save the current paragraph before switching segment type
                    self.segments.append(Segment(current_text, current_segment_type, current_importance))

                # Determine if it's a title or subtitle based on heading level
                level = int(para.style.name.split()[1])
                current_segment_type = 'title' if level == 1 else 'subtitle'
                current_importance = 2.0 if level == 1 else 1.8
                self.segments.append(Segment(line, current_segment_type, current_importance))

                # Reset for the next paragraph
                current_text = ""
                current_segment_type = "paragraph"
                current_importance = 1.0
            else:
                # Concatenate the line to the current paragraph
                current_text += line + " "

        # Add any remaining paragraph text to the segments
        if current_text:
            self.segments.append(Segment(current_text.strip(), current_segment_type, current_importance))

    def _segment_plain_text(self, text):
        """
        Segments plain text by splitting it into paragraphs.
        This method is also used for PDFs that are converted to Markdown.

        :param text: The plain text content to be segmented
        """
        paragraphs = text.split("\n")
        for paragraph in paragraphs:
            if paragraph.strip():  # Skip empty lines
                self.segments.append(Segment(paragraph.strip(), "paragraph", 1.0))

    def _save_segments_to_csv(self, output_chunk_save):
        """
        Saves the detected segments into a CSV file.

        :param output_chunk_save: The name of the output CSV file
        """
        output_csv_path = os.path.join(self.output_dir, output_chunk_save)
        with open(output_csv_path, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow(['Segment Type', 'Importance', 'Text'])
            for segment in self.segments:
                writer.writerow([segment.segment_type, segment.importance, segment.text])

    def _save_segments_to_hierarchical_json(self, output_chunk_save):
        """
        Saves the detected segments into a JSON file with a hierarchical structure (Title -> Subtitles -> Paragraphs).

        :param output_chunk_save: The name of the output JSON file
        """
        output_json_path = os.path.join(self.output_dir, output_chunk_save)

        # Initialize the JSON structure
        json_data = []
        current_title = None
        current_subtitle = None

        # Iterate through each segment and build the hierarchy
        for segment in self.segments:
            # Use match-case introduced in Python 3.10 for cleaner structure
            match segment.segment_type:
                case 'title':
                    # If it's a title, start a new section
                    current_title = {
                        "Title": segment.text,
                        "Subtitles": []  # Subtitles will be stored here
                    }
                    json_data.append(current_title)
                    current_subtitle = None  # Reset the subtitle for the new title

                case 'subtitle':
                    # If it's a subtitle, add it under the current title
                    current_subtitle = {
                        "Subtitle": segment.text,
                        "Paragraphs": []  # Paragraphs will be stored under each subtitle
                    }
                    if current_title:
                        current_title["Subtitles"].append(current_subtitle)

                case 'paragraph':
                    # If it's a paragraph, add it under the current subtitle if it exists, else under the title
                    if current_subtitle:
                        current_subtitle["Paragraphs"].append(segment.text)
                    elif current_title:
                        # If there's no subtitle, add the paragraph directly under the title
                        if "Paragraphs" not in current_title:
                            current_title["Paragraphs"] = []
                        current_title["Paragraphs"].append(segment.text)

        # Write the structured data to a JSON file
        with open(output_json_path, 'w', encoding='utf-8') as file:
            json.dump(json_data, file, ensure_ascii=False, indent=4)

        print(f"Segments have been saved in hierarchical JSON format at {output_json_path}")

    def _save_segments(self, output_chunk_name):
        """
        Generic save method to select the appropriate save method based on the format.
        :param output_chunk_name: The name of the output file (without extension).
        """
        # Define a mapping between the format and the corresponding method
        save_methods = {
            'csv': self._save_segments_to_csv,
            'json': self._save_segments_to_hierarchical_json,
        }

        # Select the appropriate save method
        save_method = save_methods.get(self.save_format)

        if save_method is None:
            raise ValueError(f"Unsupported save format: {self.save_format}")

        # Call the selected save method
        save_method(output_chunk_name)

    def process(self):
        """
        The main process for segmenting the document and saving the results into a CSV file.

        This function orchestrates the loading, segmenting, and saving of the document.
        """
        try:
            self._segment_document()

            # Generate the CSV file name based on the original file name and current timestamp
            base_filename = os.path.splitext(os.path.basename(self.file_path))[0]

            # Generate the output file name based on the original file name and current timestamp
            output_chunk_create_name = f"{base_filename}_chunk_{datetime.today().strftime('%Y_%m_%d_%H_%M_%S')}.{self.save_format}"

            # Save the segments to the appropriate format (CSV or JSON)
            self._save_segments(output_chunk_name=output_chunk_create_name)
            print(f"Segments have been saved in {os.path.join(self.output_dir, output_chunk_create_name)}")

        except ValueError as e:
            print(f"Processing error: {e}")



# Utilisation de la classe
if __name__ == "__main__":
    file_path = r'data/original/consultation_TMA_Feg_2024_Lot_3_V1.0.docx'  # Remplace par le chemin réel vers ton fichier
    output_dir = 'data/chunk'  # Chemin vers le dossier où sauvegarder le CSV
    segmenter = DocumentSegmenter(
        file_path=file_path,
        output_dir=output_dir,
        save_format='json')
    segmenter.process()